import os
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def set_cell_background(cell, color_hex):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex)
    tc_pr.append(shd)

def create_report():
    doc = Document()
    
    # Page setup
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Base Styles Configuration
    style_normal = doc.styles['Normal']
    font = style_normal.font
    font.name = 'Arial'
    font.size = Pt(11)

    # --- TITLE ---
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run("PROJECT ENGINEERING REPORT\nLOCAL-DOCUBRAIN RAG ENGINE")
    title_run.font.name = 'Arial'
    title_run.font.size = Pt(22)
    title_run.bold = True
    
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = subtitle.add_run("An Enterprise-Grade, 100% Offline Information Retrieval & Knowledge Assistant")
    sub_run.font.size = Pt(12)
    sub_run.italic = True
    
    doc.add_paragraph("-" * 80)

    # --- SECTION 1 ---
    h1 = doc.add_heading(level=1)
    h1_run = h1.add_run("1. Executive Project Summary")
    h1_run.font.name = 'Arial'
    h1_run.bold = True
    
    p1 = doc.add_paragraph(
        "The Local-DocuBrain RAG Engine project addresses a critical paradigm in modern artificial intelligence: "
        "enabling safe, private, and offline interaction with unstructured document directories. The primary objective "
        "was to construct an end-to-end Retrieval-Augmented Generation (RAG) system running entirely on local consumer hardware. "
        "By eliminating reliance on third-party cloud APIs, the system guarantees zero data leakage, complete data sovereignty, "
        "and zero operational API costs. This infrastructure functions as an 'Offline Digital Research Library,' allowing students, "
        "researchers, and developers to converse with complex text documents, extract instant summaries, and conduct cross-document analysis safely."
    )

    # --- SECTION 2 ---
    h2 = doc.add_heading(level=1)
    h2_run = h2.add_run("2. Technical Stack & Implementation Architecture")
    h2_run.font.name = 'Arial'
    h2_run.bold = True
    
    doc.add_paragraph("The underlying software architecture leverages state-of-the-art open-source components configured inside an isolated Python workspace:")
    
    techs = [
        ("Core Language", "Python 3.12 execution layer managed via isolated virtual environments (ml_env)."),
        ("Orchestration Framework", "LlamaIndex (v0.14+) utilized to construct ingestion pipelines, node segmentation, and active query interfaces."),
        ("Local Inference Engine", "Ollama background desktop abstraction layer deployed to serve open-source weights locally via port 11434."),
        ("Large Language Model", "Microsoft Phi-3 (3.8 Billion parameters) serving as the reasoning and response synthesis brain."),
        ("Local Embedding Model", "BAAI/bge-small-en-v1.5 running locally to transform text chunks into continuous mathematical vectors."),
        ("Data Parsing Layer", "PyPDF compilation library deployed to read and parse compressed document layouts into plain text formats."),
        ("Database Storage Layer", "LlamaIndex File-Based Storage Context utilized to construct and store vector indices permanently onto local disk tracks.")
    ]
    
    for title_t, desc_t in techs:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(f"{title_t}: ").bold = True
        p.add_run(desc_t)

    # --- SECTION 3 ---
    h3 = doc.add_heading(level=1)
    h3_run = h3.add_run("3. Engineering Outcomes & Key Findings")
    h3_run.font.name = 'Arial'
    h3_run.bold = True
    
    # FIX: Increased rows to 5 to accommodate header + 4 rows of data perfectly
    table = doc.add_table(rows=5, cols=2)
    table.style = 'Light Shading Accent 1'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Engineering Challenge / Milestone'
    hdr_cells[1].text = 'Observed Outcome & Empirical Finding'
    
    for cell in hdr_cells:
        set_cell_background(cell, "003366")
        cell.paragraphs[0].runs[0].font.color.rgb = None
        cell.paragraphs[0].runs[0].bold = True

    row_data = [
        ("Data Sovereignty & Privacy", "Successfully ingested 'Atomic Habits.pdf' and ran inference loops without internet access or data leaks."),
        ("Persistent Vector Caching", "Implemented disk storage context layer. First-time ingestion saves mathematical tensors to a disk folder, allowing future boots to load in under 2 seconds by skipping parsing loops."),
        ("OOM Memory Resolution", "Solved a critical 51.5 GB memory allocation crash by forcing context caps. Restricting token windows to 3072 capped background RAM overhead within stable system boundaries."),
        ("Semantic Extraction Precision", "The model successfully bypassed keyword limits. Extracted precise biographical details ('James Clear') and synthesized complex theoretical concepts through contextual vector maps.")
    ]

    for i, (col1, col2) in enumerate(row_data, start=1):
        row_cells = table.rows[i].cells
        row_cells[0].text = col1
        row_cells[1].text = col2
        if i % 2 == 0:
            set_cell_background(row_cells[0], "F2F2F2")
            set_cell_background(row_cells[1], "F2F2F2")

    doc.add_paragraph("\n")

    # --- SECTION 4 ---
    h4 = doc.add_heading(level=1)
    h4_run = h4.add_run("4. Scope of Improvements & Next Version Blueprint")
    h4_run.font.name = 'Arial'
    h4_run.bold = True
    
    improvements = [
        ("Context Window Expansion", "Transition the current 3GB Phi-3 model to Meta Llama 3 (8B) or Qwen 2 to enhance synthesis fluency and eliminate truncation loops across vast document sets."),
        ("Automated Incremental Indexing", "Incorporate directory watching utilities to cross-reference hash structures, automatically vectorizing newly added files while leaving existing caches untouched."),
        ("Graphical Web Interface Integration", "Incorporate a frontend web skin using Streamlit. This transitions the product out of the back-end terminal into a browser interface with drag-and-drop file uploads and custom user toggles.")
    ]
    
    for title_i, desc_i in improvements:
        p = doc.add_paragraph(style='List Bullet')
        p.add_run(f"{title_i}: ").bold = True
        p.add_run(desc_i)

    # Save document
    filename = "Local_DocuBrain_Project_Report.docx"
    doc.save(filename)
    print(f"🎉 Success! Professional Word Report saved as '{filename}' in your directory.")

if __name__ == "__main__":
    create_report()
